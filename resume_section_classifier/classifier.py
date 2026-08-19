"""
Resume Section Classifier — orchestrator
-------------------------------------------
Wires together: (optional) file reading -> segmentation
(rule-based heading detection + NLP fallback) -> structured storage.

Designed to slot in right after Day 5's `resume_extraction_engine`:
pass this engine the `cleaned_text` field from Day 5's
`ExtractionRecord` and it will tag every part of that text with a
section label. It can also read a `.txt` file directly, or a
`.pdf`/`.docx` file if `pdfplumber`/`python-docx` are installed
(best-effort — for production use, prefer routing files through Day
5's engine first, since it already handles multi-column PDFs, OCR
fallback, and table linearization).
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import List

from .segmenter import SectionSegmenter
from .storage import ResultStore, SectionBlockRecord, SectionClassificationRecord

logger = logging.getLogger("section_classifier")
if not logger.handlers:
    logging.basicConfig(
        level=logging.INFO, format="%(asctime)s [%(levelname)s] %(name)s: %(message)s"
    )

SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}


class UnsupportedFileTypeError(Exception):
    pass


class SectionClassifierEngine:
    def __init__(self, output_dir: str | Path = "outputs"):
        self.segmenter = SectionSegmenter()
        self.store = ResultStore(output_dir)

    # ------------------------------------------------------------------ #
    # Primary API: classify already-cleaned text (no file I/O involved).
    # This is the integration point with Day 5's ExtractionRecord.cleaned_text
    # and the recommended entry point for the matching/screening pipeline.
    # ------------------------------------------------------------------ #
    def classify_text(self, text: str, source_name: str = "input") -> SectionClassificationRecord:
        warnings: List[str] = []

        if not text or not text.strip():
            return SectionClassificationRecord(
                source_file=source_name,
                classified_at=ResultStore.now_iso(),
                total_char_count=0,
                sections_detected=[],
                heading_based_blocks=0,
                nlp_based_blocks=0,
                uncategorized_char_count=0,
                blocks=[],
                warnings=["Input text was empty."],
                status="failed",
                error="Empty input text.",
            )

        raw_blocks = self.segmenter.segment(text)

        block_records: List[SectionBlockRecord] = []
        heading_based = 0
        nlp_based = 0
        uncategorized_chars = 0

        for b in raw_blocks:
            if b.method == "rule_heading":
                heading_based += 1
            else:
                nlp_based += 1
            if b.label == "Uncategorized":
                uncategorized_chars += len(b.text)
            if b.flags:
                warnings.extend(b.flags)

            block_records.append(
                SectionBlockRecord(
                    label=b.label,
                    method=b.method,
                    confidence=b.confidence,
                    start_line=b.start_line,
                    end_line=b.end_line,
                    char_count=len(b.text),
                    text=b.text,
                    heading_text=b.heading_text,
                    nlp_label=b.nlp_label,
                    nlp_confidence=b.nlp_confidence,
                    flags=b.flags,
                )
            )

        sections_detected = []
        for b in block_records:
            if b.label not in sections_detected:
                sections_detected.append(b.label)

        status = "success"
        if uncategorized_chars > 0.4 * len(text):
            status = "partial"
            warnings.append("More than 40% of the document could not be confidently classified.")

        return SectionClassificationRecord(
            source_file=source_name,
            classified_at=ResultStore.now_iso(),
            total_char_count=len(text),
            sections_detected=sections_detected,
            heading_based_blocks=heading_based,
            nlp_based_blocks=nlp_based,
            uncategorized_char_count=uncategorized_chars,
            blocks=block_records,
            warnings=warnings,
            status=status,
        )

    # ------------------------------------------------------------------ #
    # File-based API
    # ------------------------------------------------------------------ #
    def process_file(self, path: str | Path) -> SectionClassificationRecord:
        path = Path(path)
        ext = path.suffix.lower()

        if ext not in SUPPORTED_EXTENSIONS:
            raise UnsupportedFileTypeError(
                f"Unsupported file type '{ext}'. Supported: {sorted(SUPPORTED_EXTENSIONS)} "
                "(for .pdf/.docx with full layout handling, run resume_extraction_engine "
                "first and pass its cleaned_text to classify_text())."
            )
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        logger.info("Classifying sections for %s", path.name)

        try:
            text = self._read_text(path, ext)
            record = self.classify_text(text, source_name=str(path))
        except Exception as exc:  # noqa: BLE001
            logger.exception("Failed to classify %s", path)
            record = SectionClassificationRecord(
                source_file=str(path),
                classified_at=ResultStore.now_iso(),
                total_char_count=0,
                sections_detected=[],
                heading_based_blocks=0,
                nlp_based_blocks=0,
                uncategorized_char_count=0,
                blocks=[],
                warnings=[],
                status="failed",
                error=str(exc),
            )

        self.store.save(record)
        return record

    def process_directory(self, directory: str | Path) -> List[SectionClassificationRecord]:
        directory = Path(directory)
        records = []
        for file in sorted(directory.iterdir()):
            if file.suffix.lower() in SUPPORTED_EXTENSIONS:
                records.append(self.process_file(file))
        return records

    @staticmethod
    def _read_text(path: Path, ext: str) -> str:
        if ext == ".txt":
            return path.read_text(encoding="utf-8")

        if ext == ".docx":
            try:
                import docx
            except ImportError as exc:
                raise RuntimeError(
                    "python-docx is not installed; install it or pre-clean the file with "
                    "resume_extraction_engine and call classify_text() instead."
                ) from exc
            document = docx.Document(str(path))
            lines = [p.text for p in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    lines.append(" | ".join(cell.text.strip() for cell in row.cells))
            return "\n".join(lines)

        if ext == ".pdf":
            try:
                import pdfplumber
            except ImportError as exc:
                raise RuntimeError(
                    "pdfplumber is not installed; install it or pre-clean the file with "
                    "resume_extraction_engine and call classify_text() instead."
                ) from exc
            chunks = []
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    chunks.append(page.extract_text() or "")
            return "\n".join(chunks)

        raise UnsupportedFileTypeError(ext)
